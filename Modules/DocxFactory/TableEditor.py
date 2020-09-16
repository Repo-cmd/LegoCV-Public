from docx import Document

def TableEditor(TargetDocument, Keyword, Value, OutputFile):

    """
    **TargetDocument: str of the location and file name

    **Keyword: list of the keyword that will be replaced

    **Value:  list of the desired value

    **OutputFile: str of the location and file name (replaces files with the same name)

    Searches in a document in a directory to replace a keyword with a value using the same style

    P.s. For tables only (not free text) and copies style only
    """
    document = Document(str(TargetDocument))
    # loop through each table
    for table in document.tables:
        # loop through each row of the current table
        for row in table.rows:
            # loop through each cell of the current row
            for cell in row.cells:
                for key in Keyword:
                    # check if the key exists
                    if str(key) in cell.text:
                        # copy current style
                        styleHolder = cell.paragraphs[0].style
                        x = cell.text
                        # replace the key with the inputted value and apply old style
                        cell.text = x.replace(str(key), str(Value[Keyword.index(key)]))
                        cell.paragraphs[0].style = styleHolder
                        #remove the values and the keys used
                        Value.pop(Keyword.index(key))
                        Keyword.pop(Keyword.index(key))
    # save the document in the provided directory with the provided name
    document.save(str(OutputFile))