import docx

def TableDuplicator(TargetDocument, RepCount, OutputFile, RepMarker):

    """
    **TargetDocument: str of the location and file name

    **RepCount: a list of integers that represent the number of repetition per table block

    **OutputFile: str of the location and file name (replaces files with the same name)

    **RepMarker: a list of strings that marks a table for repetition

    searches in a document in a directory for a RepMarker to repeat it n times
    
    [Creates an output file]

    P.s. For tables only (not free text) and copies style only
    """
    document = docx.Document(str(TargetDocument))
    
    # loop through each table
    for table in document.tables:
        mark = ''
        # loop through each row of the current table for a mark
        for row in table.rows:
            # loop through each cell of the current row
            for cell in row.cells:
                for repMark in RepMarker:
                    # check if the key exists
                    if str(repMark) in cell.text:
                       mark = 'yes'
                       CurrentValue = repMark
        
        # copy current rows text and style
        if mark == 'yes' and int(RepCount[RepMarker.index(CurrentValue)]) > 0:
            cellStyle =[]
            cellTxt = []
            for cell in table._cells:
                cellStyle.append(cell.paragraphs[0].style)
                cellTxt.append(cell.text)
            cellStyle = cellStyle * int(int(RepCount[RepMarker.index(CurrentValue)]) + 1) 
            cellTxt = cellTxt * int(int(RepCount[RepMarker.index(CurrentValue)]) + 1) 
            
            # table rows duplication
            for _ in range(int(len(table.rows)) * int(RepCount[RepMarker.index(CurrentValue)])):
                table.add_row()
            
            # duplicate text and styles
            x = range(len(table._cells))
            
            for item in x:
                table._cells[x.index(item)].text = cellTxt[x.index(item)]
                table._cells[x.index(item)].paragraphs[0].style = cellStyle[x.index(item)]

    # save the document in the provided directory with the provided name
    document.save(str(OutputFile))